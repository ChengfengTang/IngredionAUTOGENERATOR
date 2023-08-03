import pandas as pd
from docx import Document
from copy import deepcopy
from datetime import datetime
import tkinter as tk
from tkinter import filedialog, messagebox, Checkbutton, IntVar, Label

class App:
    def __init__(self, master):
        self.master = master
        self.PO_check_vars = []
        self.excel_file_path = ""
        self.load_button = tk.Button(master, text='Load Excel', command=self.load_excel)
        self.load_button.pack()
        self.template1_entry = self.create_file_entry("Select COD Template File")
        self.template2_entry = self.create_file_entry("Select Email Template File")
        self.output_entry = self.create_directory_entry("Select Output Directory")
        self.generate_button = tk.Button(master, text="Generate", command=self.generate)
        self.generate_button.pack()

    def create_file_entry(self, button_text):
        entry = tk.Entry(self.master, width=50)
        entry.pack()
        button = tk.Button(self.master, text=button_text, command=lambda: self.select_file(entry))
        button.pack()
        return entry

    def create_directory_entry(self, button_text):
        entry = tk.Entry(self.master, width=50)
        entry.pack()
        button = tk.Button(self.master, text=button_text, command=lambda: self.select_directory(entry))
        button.pack()
        return entry

    def load_excel(self):
        self.excel_file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        df = pd.read_excel(self.excel_file_path)
        # Filter to keep only integer PO numbers
        integer_PO_numbers = [PO for PO in df['PO'].unique() if str(PO).isdigit()]
        self.create_checkbuttons(integer_PO_numbers)

    def select_file(self, entry):
        file_path = filedialog.askopenfilename()
        entry.delete(0, tk.END)
        entry.insert(0, file_path)

    def select_directory(self, entry):
        directory_path = filedialog.askdirectory()
        entry.delete(0, tk.END)
        entry.insert(0, directory_path)

    def create_checkbuttons(self, PO_numbers):
        for PO in PO_numbers:
            var = IntVar()
            chk = Checkbutton(self.master, text=str(PO), variable=var)
            chk.pack()
            self.PO_check_vars.append(var)

    def generate(self):
        try:
            # Load the Excel file
            data = pd.read_excel(self.excel_file_path)

            # Convert 'PO' to numeric and drop rows with non-numeric 'PO'
            data['PO'] = pd.to_numeric(data['PO'], errors='coerce')
            data = data.dropna(subset=['PO'])

            # Only keep rows with the user-specified PO numbers
            selected_PO_numbers = [PO for PO, var in zip(data['PO'].unique(), self.PO_check_vars) if var.get()]
            data = data[data['PO'].isin(selected_PO_numbers)]

            # Load the Word documents
            template1 = Document(self.template1_entry.get())
            template2 = Document(self.template2_entry.get())

            # Get today's date
            today = datetime.today().strftime('%Y/%m/%d')

            # Group the DataFrame by the 'PO' column
            groups = data.groupby('PO')

            for po, group in groups:
                # Create a copy of the templates for each group
                doc1 = deepcopy(template1)
                doc2 = deepcopy(template2)

                # Combine the quantities, batch numbers, and reasons in the desired format
                damages_reasons = ', '.join(
                    f'{int(row["Damage"])} bags {row["Reason"]} ({row["Batch Number"]})'
                    for _, row in group.iterrows()
                    if row['Reason'] not in ['Shortage', 'Extra']
                )

                # Combine the quantities and reasons for the second document
                group2 = group.groupby('Reason').agg({'Damage': 'sum'}).reset_index()
                damages_reasons2 = ', and '.join(
                    f'{int(row["Damage"])} bags are found {row["Reason"]}'
                    for _, row in group2.iterrows()
                )

                filtered_group = group[group['Reason'].isin(['Shortage', 'Extra']) == False]

                # Define a dictionary that maps placeholders to replacement text
                replacements1 = {
                    '<DATE>': today,
                    '<PO>': str(int(po)),  # convert 'PO' to int before converting to string to remove decimal point
                    '<INVOICE_NO>': ', '.join(set(group['Invoice NO'].astype(int).astype(str))),
                    '<PRODUCT_CODE>': ', '.join(set(group['Product code'].astype(str))),
                    '<DAMAGE>': damages_reasons,
                    '<BATCH_NUMBER>': ' / '.join(filtered_group['Batch Number'].astype(str)),
                    '<REASON>': ', '.join(set(filtered_group['Reason'].astype(str))),
                }

                replacements2 = {
                    '<REASON>': damages_reasons2,
                    '<COMPLAINT>': ', '.join(set(group['Complaint'].astype(str))),
                    '<PO>': str(int(po)),
                    '<INVOICE_NUM>': ', '.join(set(group['Invoice NO'].astype(int).astype(str))),
                    '<BAG_NUM>': str(int(filtered_group['Damage'].sum())),
                    '<VALUE>': str(round(group['Good Value'].sum(), 2)),
                }

                # Replace the placeholders in the documents
                self.replace_placeholders(doc1, replacements1)
                self.replace_placeholders(doc2, replacements2)

                for table in doc2.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            self.replace_placeholders_in_cell(cell, replacements2)

                # Save the changes with a filename based on the 'PO' value
                doc1.save(f'{self.output_entry.get()}/COD {int(po)}.docx')  # convert 'PO' to int to remove decimal point
                doc2.save(f'{self.output_entry.get()}/Email {int(po)}.docx')  # convert 'PO' to int to remove decimal point

            # Show a success message and exit the program
            messagebox.showinfo("Success", "Generation completed successfully!")
            self.master.quit()
        except Exception as e:
            # Show an error message if anything goes wrong
            messagebox.showerror("Error", str(e))

    def replace_placeholders(self, doc, replacements):
        for para in doc.paragraphs:
            for placeholder, replacement_text in replacements.items():
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, replacement_text)

    def replace_placeholders_in_cell(self, cell, replacements):
        for para in cell.paragraphs:
            for placeholder, replacement_text in replacements.items():
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, replacement_text)

root = tk.Tk()
app = App(root)
root.mainloop()
